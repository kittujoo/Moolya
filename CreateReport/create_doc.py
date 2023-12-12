from docx import Document

def create_docx(file_path):
    # Create a new Word document
    doc = Document()

    # Add some text to the document
    doc.add_heading('My Python-Generated Document', level=1)

    doc.add_paragraph("This is a paragraph of text in the Word document. ")
    doc.add_paragraph("You can add more paragraphs as needed.")
    doc.add_paragraph(" ಕನ್ನಡದಲ್ಲಿ ಪೈಥಾನ್ ತರಬೇತಿ")

    # Save the document
    doc.save(file_path)

# Example usage:
docx_file_path = "example_document1.docx"
create_docx(docx_file_path)

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement

def add_kannada_text(doc, kannada_text, font_path):
    # Add Kannada text to the document
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    run.text = kannada_text

    # Set Kannada font for the run
    font_element = run._element
    font = OxmlElement('w:rFonts')
    font.set(qn('w:ascii'), 'Lohit Kannada')
    font.set(qn('w:hAnsi'), 'Lohit Kannada')
    font.set(qn('w:cs'), 'Lohit Kannada')
    font_element.insert(0, font)

    # Adjust font size if needed
    run.font.size = Pt(12)

# Example usage:
doc = Document()
kannada_text = "ಹಲೋ, ಈ ಒಂದು ಕನ್ನಡ ಡಾಕ್ಯುಮೆಂಟ್ ಆಗಿದೆ!"

add_kannada_text(doc, kannada_text, 'Lohit-Kannada.ttf')

doc.save("kannada_document.docx")

